#!/usr/bin/env python3
"""
md2docx.py — Markdown to DOCX converter
支持：标题、段落、粗体/斜体/删除线/行内代码、有序/无序列表（多级）、
      表格、代码块、引用块、水平线、图片（本地 & 网络）、超链接
用法：
    python md2docx.py input.md output.docx
    python md2docx.py input.md              # 输出 input.docx
    python md2docx.py *.md -o out_dir/      # 批量转换
"""

import argparse
import base64
import io
import os
import re
import sys
from pathlib import Path
from urllib.request import urlopen, Request
from urllib.error import URLError

try:
    import mistune
except ImportError:
    sys.exit("请先安装 mistune：pip install mistune")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    sys.exit("请先安装 python-docx：pip install python-docx")

try:
    from PIL import Image as PILImage

    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# ── 颜色 & 字号常量 ──────────────────────────────────────────────
HEADING_SIZES = {1: 28, 2: 24, 3: 20, 4: 18, 5: 16, 6: 14}
HEADING_COLORS = {
    1: RGBColor(0x1F, 0x49, 0x7D),
    2: RGBColor(0x2E, 0x74, 0xB5),
    3: RGBColor(0x2E, 0x74, 0xB5),
    4: RGBColor(0x4A, 0x4A, 0x4A),
    5: RGBColor(0x4A, 0x4A, 0x4A),
    6: RGBColor(0x6A, 0x6A, 0x6A),
}
CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)
CODE_FONT = "Courier New"
QUOTE_COLOR = RGBColor(0x60, 0x60, 0x60)
BODY_FONT = "Arial"


# ── 工具函数 ─────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """给表格单元格设置背景色（hex，如 'D9E1F2'）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_paragraph_border(paragraph, side="left", color="2E74B5", size=24, space=4):
    """给段落加左侧边框，用于 blockquote"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), str(size))
    bdr.set(qn("w:space"), str(space))
    bdr.set(qn("w:color"), color)
    pBdr.append(bdr)
    pPr.append(pBdr)


def add_horizontal_rule(doc):
    """插入水平线（段落下边框）"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_hyperlink(paragraph, url: str, text: str):
    """在段落中插入超链接"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def load_image(src: str, base_dir: str):
    """加载图片，返回 (bytes, extension)，失败返回 None"""
    try:
        if src.startswith("data:"):
            m = re.match(r"data:image/(\w+);base64,(.+)", src)
            if m:
                ext = m.group(1).lower()
                return base64.b64decode(m.group(2)), ext
        elif src.startswith(("http://", "https://")):
            req = Request(src, headers={"User-Agent": "Mozilla/5.0"})
            data = urlopen(req, timeout=10).read()
            ext = src.split(".")[-1].split("?")[0].lower() or "png"
            return data, ext
        else:
            path = Path(base_dir) / src if not Path(src).is_absolute() else Path(src)
            if path.exists():
                return path.read_bytes(), path.suffix.lstrip(".").lower() or "png"
    except Exception as e:
        print(f"  ⚠ 图片加载失败 [{src}]: {e}", file=sys.stderr)
    return None, None


def image_dimensions(img_bytes: bytes, max_width_inches: float = 6.0):
    """根据图片原始尺寸计算合适的插入宽高（保持比例，最大宽度限制）"""
    if not HAS_PIL:
        return Inches(max_width_inches), None
    try:
        img = PILImage.open(io.BytesIO(img_bytes))
        w_px, h_px = img.size
        dpi = img.info.get("dpi", (96, 96))
        dpi_x = dpi[0] if isinstance(dpi, tuple) else dpi
        w_in = w_px / dpi_x
        if w_in > max_width_inches:
            ratio = max_width_inches / w_in
            w_in = max_width_inches
            h_in = (h_px / dpi_x) * ratio
        else:
            h_in = h_px / dpi_x
        return Inches(w_in), Inches(h_in)
    except Exception:
        return Inches(max_width_inches), None


# ── 行内元素解析 ─────────────────────────────────────────────────

# 行内标记解析：粗体、斜体、粗斜体、行内代码、删除线、超链接、图片
INLINE_PATTERNS = [
    ("img", re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')),
    ("link", re.compile(r'\[([^\]]+)\]\(([^)]+)\)')),
    ("bold_it", re.compile(r'\*{3}(.+?)\*{3}|_{3}(.+?)_{3}')),
    ("bold", re.compile(r'\*{2}(.+?)\*{2}|_{2}(.+?)_{2}')),
    ("italic", re.compile(r'\*(.+?)\*|_(.+?)_')),
    ("strike", re.compile(r'~~(.+?)~~')),
    ("code", re.compile(r'`(.+?)`')),
]


def parse_inline(paragraph, text: str, base_dir: str, defaults: dict = None):
    """
    解析行内 Markdown 并向 paragraph 追加 runs。
    defaults: {'bold': bool, 'italic': bool, 'color': RGBColor, 'font': str, 'size': int}
    """
    defaults = defaults or {}

    def apply_defaults(run):
        run.font.name = defaults.get("font", BODY_FONT)
        if "size" in defaults:
            run.font.size = Pt(defaults["size"])
        if defaults.get("bold"):
            run.bold = True
        if defaults.get("italic"):
            run.italic = True
        if "color" in defaults:
            run.font.color.rgb = defaults["color"]

    def add_run(para, txt, bold=False, italic=False, code=False, strike=False, color=None):
        if not txt:
            return
        run = para.add_run(txt)
        apply_defaults(run)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if strike:
            run.font.strike = True
        if code:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
            # 行内代码底纹（shading）
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F0F0F0")
            rPr.append(shd)
        if color:
            run.font.color.rgb = color

    # 递归切分 token
    def process(para, segment: str, **fmt):
        if not segment:
            return
        # 找最先出现的模式
        earliest = None
        for name, pat in INLINE_PATTERNS:
            m = pat.search(segment)
            if m and (earliest is None or m.start() < earliest[1].start()):
                earliest = (name, m)
        if not earliest:
            add_run(para, segment, **fmt)
            return
        name, m = earliest
        # 前缀
        if m.start() > 0:
            add_run(para, segment[: m.start()], **fmt)
        # 匹配部分
        if name == "img":
            alt, src = m.group(1), m.group(2)
            img_bytes, ext = load_image(src, base_dir)
            if img_bytes:
                width, height = image_dimensions(img_bytes)
                try:
                    stream = io.BytesIO(img_bytes)
                    kw = {"width": width}
                    if height:
                        kw["height"] = height
                    run = para.add_run()
                    run.add_picture(stream, **kw)
                except Exception as e:
                    add_run(para, f"[图片: {alt}]", color=RGBColor(0xAA, 0x00, 0x00))
                    print(f"  ⚠ 图片插入失败: {e}", file=sys.stderr)
            else:
                add_run(para, f"[图片: {alt}]", color=RGBColor(0xAA, 0x00, 0x00))
        elif name == "link":
            link_text, url = m.group(1), m.group(2)
            add_hyperlink(para, url, link_text)
        elif name == "bold_it":
            inner = m.group(1) or m.group(2)
            process(para, inner, bold=True, italic=True)
        elif name == "bold":
            inner = m.group(1) or m.group(2)
            process(para, inner, bold=True)
        elif name == "italic":
            inner = m.group(1) or m.group(2)
            process(para, inner, italic=True)
        elif name == "strike":
            process(para, m.group(1), strike=True)
        elif name == "code":
            add_run(para, m.group(1), code=True)
        # 后缀
        process(para, segment[m.end():], **fmt)

    process(paragraph, text)


# ── 主转换器 ─────────────────────────────────────────────────────

class MarkdownToDocx:
    def __init__(self, md_path: str, docx_path: str):
        self.md_path = Path(md_path)
        self.docx_path = Path(docx_path)
        self.base_dir = str(self.md_path.parent)
        self.doc = Document()
        self._setup_styles()
        # 列表深度追踪
        self._list_stack: list[dict] = []  # [{"type": "ul"|"ol", "count": int}]

    # ── 样式初始化 ────────────────────────────────────────────────
    def _setup_styles(self):
        """设置文档默认字体、段落间距"""
        style = self.doc.styles["Normal"]
        style.font.name = BODY_FONT
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        # 确保内置 Heading 样式颜色正确
        for level, size in HEADING_SIZES.items():
            h_name = f"Heading {level}"
            if h_name in self.doc.styles:
                hs = self.doc.styles[h_name]
                hs.font.name = BODY_FONT
                hs.font.size = Pt(size)
                hs.font.color.rgb = HEADING_COLORS[level]
                hs.font.bold = True
                hs.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
                hs.paragraph_format.space_after = Pt(6)

    # ── 解析入口 ──────────────────────────────────────────────────
    def convert(self):
        text = self.md_path.read_text(encoding="utf-8")
        self._parse(text)
        self.doc.save(str(self.docx_path))
        print(f"✅ 转换完成：{self.docx_path}")

    # ── 文本预处理 ────────────────────────────────────────────────
    def _parse(self, text: str):
        """逐块解析 Markdown"""
        # 统一换行
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # —— 代码块 ——
            if line.startswith("```"):
                lang = line[3:].strip()
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                self._add_code_block("\n".join(code_lines), lang)
                i += 1
                continue

            # —— 水平线 ——
            if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", line):
                add_horizontal_rule(self.doc)
                i += 1
                continue

            # —— 标题 ——
            m = re.match(r"^(#{1,6})\s+(.+)", line)
            if m:
                level = len(m.group(1))
                self._add_heading(m.group(2), level)
                self._list_stack.clear()
                i += 1
                continue

            # —— 表格（收集连续行）——
            if "|" in line and i + 1 < len(lines) and re.match(r"^\|?[\s\-|:]+\|?$", lines[i + 1]):
                table_lines = [line]
                i += 1
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(table_lines)
                self._list_stack.clear()
                continue

            # —— 引用块 ——
            if line.startswith(">"):
                quote_lines = []
                while i < len(lines) and lines[i].startswith(">"):
                    quote_lines.append(lines[i].lstrip(">").strip())
                    i += 1
                self._add_blockquote(" ".join(quote_lines))
                continue

            # —— 无序列表 ——
            m = re.match(r"^(\s*)([-*+])\s+(.+)", line)
            if m:
                indent = len(m.group(1))
                level = indent // 2
                self._add_list_item(m.group(3), level, ordered=False)
                i += 1
                continue

            # —— 有序列表 ——
            m = re.match(r"^(\s*)(\d+)[.)]\s+(.+)", line)
            if m:
                indent = len(m.group(1))
                level = indent // 2
                self._add_list_item(m.group(3), level, ordered=True)
                i += 1
                continue

            # —— 空行 ——
            if line.strip() == "":
                self._list_stack.clear()
                i += 1
                continue

            # —— 普通段落 ——
            self._list_stack.clear()
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", ">", "```", "|")):
                if re.match(r"^(\s*)([-*+]|\d+[.)])\s+", lines[i]):
                    break
                if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", lines[i]):
                    break
                para_lines.append(lines[i])
                i += 1
            self._add_paragraph(" ".join(para_lines))

    # ── 块级元素添加方法 ─────────────────────────────────────────

    def _add_heading(self, text: str, level: int):
        p = self.doc.add_heading("", level=level)
        p.clear()
        # 去除 ATX 结尾的 #
        text = re.sub(r"\s+#+\s*$", "", text).strip()
        parse_inline(p, text, self.base_dir, defaults={
            "bold": True,
            "size": HEADING_SIZES.get(level, 14),
            "color": HEADING_COLORS.get(level, RGBColor(0, 0, 0)),
            "font": BODY_FONT,
        })
        p.style = self.doc.styles[f"Heading {level}"]

    def _add_paragraph(self, text: str):
        p = self.doc.add_paragraph()
        parse_inline(p, text, self.base_dir)

    def _add_blockquote(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_paragraph_border(p, "left", color="2E74B5", size=24, space=4)
        parse_inline(p, text, self.base_dir, defaults={"color": QUOTE_COLOR, "italic": True})

    def _add_list_item(self, text: str, level: int, ordered: bool):
        """添加列表项，支持多级缩进"""
        # 计算此 level 下的序号（有序列表）
        # 调整 stack 深度
        while len(self._list_stack) > level + 1:
            self._list_stack.pop()
        while len(self._list_stack) < level + 1:
            self._list_stack.append({"type": "ol" if ordered else "ul", "count": 0})
        entry = self._list_stack[level]
        if ordered:
            entry["count"] += 1
            num = entry["count"]
            prefix = "  " * level + f"{num}. "
        else:
            bullets = ["•", "◦", "▪"]
            prefix = "  " * level + bullets[level % 3] + " "

        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(prefix)
        run.font.name = BODY_FONT
        parse_inline(p, text, self.base_dir)

    def _add_code_block(self, code: str, lang: str = ""):
        """灰色背景的代码块"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # 段落底纹
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)
        # 边框
        pBdr = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            bdr = OxmlElement(f"w:{side}")
            bdr.set(qn("w:val"), "single")
            bdr.set(qn("w:sz"), "4")
            bdr.set(qn("w:space"), "4")
            bdr.set(qn("w:color"), "CCCCCC")
            pBdr.append(bdr)
        pPr.append(pBdr)
        run = p.add_run(code)
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)

    def _add_table(self, lines: list[str]):
        """解析 Markdown 表格并插入 docx 表格"""

        # 分离表头、分隔行、数据行
        def split_row(line):
            line = line.strip().strip("|")
            return [c.strip() for c in line.split("|")]

        rows = [split_row(l) for l in lines]
        # 第 2 行是分隔行，跳过
        if len(rows) < 2:
            return
        header = rows[0]
        data_rows = rows[2:] if len(rows) > 2 else []
        col_count = len(header)

        # 对齐方式（从分隔行解析）
        alignments = []
        if len(rows) > 1:
            for cell in rows[1]:
                cell = cell.strip()
                if cell.startswith(":") and cell.endswith(":"):
                    alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                elif cell.endswith(":"):
                    alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                else:
                    alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
        while len(alignments) < col_count:
            alignments.append(WD_ALIGN_PARAGRAPH.LEFT)

        table = self.doc.add_table(rows=1 + len(data_rows), cols=col_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        def fill_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = align
            parse_inline(p, text, self.base_dir, defaults={"bold": bold})
            if bg:
                set_cell_bg(cell, bg)

        # 表头
        hdr_row = table.rows[0]
        for j, hdr in enumerate(header[:col_count]):
            fill_cell(hdr_row.cells[j], hdr, bold=True, align=alignments[j], bg="D9E1F2")

        # 数据行
        for i, row in enumerate(data_rows):
            tr = table.rows[i + 1]
            for j in range(col_count):
                cell_text = row[j] if j < len(row) else ""
                bg = "F2F2F2" if i % 2 == 1 else None
                fill_cell(tr.cells[j], cell_text, align=alignments[j], bg=bg)

        # 表格后加空段
        self.doc.add_paragraph()


# ── CLI ──────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="将 Markdown 文件转换为 DOCX（支持图片、表格、列表、代码块等）"
    )
    parser.add_argument("inputs", nargs="+", help=".md 文件路径（支持多个）")
    parser.add_argument("-o", "--output", default=None,
                        help="输出文件或目录（单文件转换可直接指定 .docx 路径）")
    args = parser.parse_args()

    inputs = [Path(p) for p in args.inputs]
    out = Path(args.output) if args.output else None

    for md_path in inputs:
        if not md_path.exists():
            print(f"❌ 文件不存在：{md_path}", file=sys.stderr)
            continue
        if out is None or (len(inputs) == 1 and out and out.suffix == ".docx"):
            docx_path = out if out else md_path.with_suffix(".docx")
        else:
            # 输出到目录
            out.mkdir(parents=True, exist_ok=True)
            docx_path = out / md_path.with_suffix(".docx").name

        print(f"🔄 转换中：{md_path} → {docx_path}")
        try:
            converter = MarkdownToDocx(str(md_path), str(docx_path))
            converter.convert()
        except Exception as e:
            print(f"❌ 转换失败：{e}", file=sys.stderr)
            import traceback
            traceback.print_exc()


if __name__ == "__main__":
    main()
